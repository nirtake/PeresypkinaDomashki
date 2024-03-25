from docx import Document
from pdfminer.high_level import extract_text
import argparse

parser = argparse.ArgumentParser()
parser.add_argument("-p", "--path", help="Введите путь к файлу, -p или --path=", required=True, type=str)
args = parser.parse_args()

def pdf_to_txt(path):
    try:
        text = extract_text(f'{path}.')
        text_file = open('converted_pdf.txt', 'w', encoding='utf8')
        print(text, file=text_file)
        print('Файл успешнео конвертирован')
    except FileNotFoundError:
        print('Файл не найден', 'red')
    return text_file

def docx_to_txt(path):
    try:
        doc = Document(f'{path}')
        text_file = open('converted_txt.txt', 'w', encoding='utf8')
        for x in range (0, len(doc.paragraphs)):
            print(doc.paragraphs[x].text, file=text_file)
        text_file.close()
        print('Файл успешнео конвертирован')
        return text_file
    except FileNotFoundError:
        print('Файл не найден')
    return text_file

def file_to_txt(path):
    if path.endswith('.docx'):
        docx_to_txt(path)
    elif path.endswith('.pdf'):
        pdf_to_txt(path)
    else:
        print('Неверный формат файла')
    


file_to_txt(args.path)


